"""
Streamlit 웹앱 → Word(.docx) 변환 스크립트

사용법:
    python scripts/export_docx.py

PDF 레이아웃(발표문.pdf)을 참고하여 동일한 배치로 Word 문서를 생성합니다.
"""

import os
import sys
import re
import json
import html as html_module
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import pandas as pd
import plotly.graph_objects as go

# ============================================================
# 경로 설정
# ============================================================
PROJECT_ROOT = Path(__file__).resolve().parent.parent
SECTIONS_DIR = PROJECT_ROOT / "app" / "sections"
TABLES_PATH = PROJECT_ROOT / "app" / "tables.json"
IMAGES_DIR = PROJECT_ROOT / "images"
OUTPUT_PATH = PROJECT_ROOT / "발표문.docx"
PDF_PATH = PROJECT_ROOT / "발표문.pdf"
TEMP_DIR = PROJECT_ROOT / "scripts" / "_temp_docx"

# ============================================================
# 스타일 상수 (PDF 레이아웃 기준)
# ============================================================
FONT_KR = "맑은 고딕"
FONT_EN = "Times New Roman"
FONT_MONO = "Consolas"

BODY_SIZE = Pt(12)
H1_SIZE = Pt(14)       # 장 제목 (Ⅰ, Ⅱ, ...)
H2_SIZE = Pt(13)       # 절 제목 (1., 2., ...)
H3_SIZE = Pt(12)       # 항 제목 (소거 사례 ①, ...)
TITLE_SIZE = Pt(20)
SUBTITLE_SIZE = Pt(14)
AUTHOR_SIZE = Pt(12)
FOOTNOTE_SIZE = Pt(9)
CAPTION_SIZE = Pt(9)
TABLE_SIZE = Pt(10)
TOC_MAIN_SIZE = Pt(11)
TOC_SUB_SIZE = Pt(10)
BLOCKQUOTE_SIZE = Pt(11)
PAGE_NUM_SIZE = Pt(10)
WEB_URL_SIZE = Pt(8)

GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
BLUE_LINK = RGBColor(0x00, 0x68, 0xC9)
TABLE_HEADER_BG = "D9E2F3"
TABLE_ALT_BG = "F2F2F2"

WEB_URL = "https://phenomenal-realism-kr-8gr9edsgfre4exrwnbbqrq.streamlit.app/"


# ============================================================
# 유틸리티 함수
# ============================================================

def set_run_font(run, font_name=None, size=None, bold=None, italic=None,
                 color=None, superscript=None):
    """Run의 폰트를 설정 (동아시아 폰트 포함)"""
    fn = font_name or FONT_KR
    run.font.name = fn
    # 동아시아 폰트 설정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:eastAsia="{fn}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), fn)
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = color
    if superscript is not None:
        run.font.superscript = superscript


def set_paragraph_format(para, space_before=None, space_after=None,
                         line_spacing=None, alignment=None, indent_left=None,
                         first_line_indent=None, keep_with_next=None):
    """문단 서식 설정"""
    fmt = para.paragraph_format
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if alignment is not None:
        fmt.alignment = alignment
    if indent_left is not None:
        fmt.left_indent = indent_left
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next


def clean_html_entities(text):
    """HTML 엔티티 디코딩 + <표 N> / <그림 N> 보호"""
    text = text.replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&amp;', '&').replace('&nbsp;', ' ')
    # <표 N>, <그림 N> → 꺾쇠가 HTML 태그로 인식되지 않도록 전각 괄호로 변환
    text = re.sub(r'<(표\s*\d+)>', r'〈\1〉', text)
    text = re.sub(r'<(그림\s*\d+)>', r'〈\1〉', text)
    return text


def strip_html_tags(text):
    """HTML 태그 제거 (태그 내 텍스트는 유지)"""
    return re.sub(r'<[^>]+>', '', text)


def add_formatted_runs(para, text, base_size=None, base_color=None):
    """마크다운/HTML 인라인 서식을 파싱하여 Run으로 추가"""
    if not text:
        return

    text = clean_html_entities(text)
    sz = base_size or BODY_SIZE
    clr = base_color

    # 인라인 패턴 매칭
    pattern = (
        r'(\*\*[^*]+?\*\*'                          # **bold**
        r'|<sup>[^<]+?</sup>'                        # <sup>text</sup>
        r'|<b>[^<]+?</b>'                            # <b>text</b>
        r'|<code>[^<]+?</code>'                      # <code>text</code>
        r'|<i>[^<]+?</i>'                            # <i>text</i>
        r'|\*[^*]+?\*'                               # *italic*
        r'|<a\s+href="[^"]*"[^>]*>[^<]*</a>'        # <a href>text</a>
        r'|<br\s*/?>)'                               # <br>
    )

    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=sz, bold=True, color=clr)
        elif part.startswith('<sup>') and part.endswith('</sup>'):
            inner = re.sub(r'</?sup>', '', part)
            run = para.add_run(inner)
            set_run_font(run, size=Pt(max(6, sz.pt - 3)), superscript=True, color=clr)
        elif part.startswith('<b>') and part.endswith('</b>'):
            inner = re.sub(r'</?b>', '', part)
            run = para.add_run(inner)
            set_run_font(run, size=sz, bold=True, color=clr)
        elif part.startswith('<code>') and part.endswith('</code>'):
            inner = re.sub(r'</?code>', '', part)
            run = para.add_run(inner)
            set_run_font(run, font_name=FONT_MONO, size=Pt(sz.pt - 1), color=clr)
        elif part.startswith('<i>') and part.endswith('</i>'):
            inner = re.sub(r'</?i>', '', part)
            run = para.add_run(inner)
            set_run_font(run, size=sz, italic=True, color=clr)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            set_run_font(run, size=sz, italic=True, color=clr)
        elif part.startswith('<a '):
            match = re.match(r'<a\s+href="([^"]*)"[^>]*>([^<]*)</a>', part)
            if match:
                _, link_text = match.groups()
                run = para.add_run(link_text)
                set_run_font(run, size=sz, color=BLUE_LINK)
        elif '<br' in part:
            run = para.add_run('\n')
        else:
            # 남은 HTML 태그 제거
            clean = re.sub(r'<[^>]+>', '', part)
            if clean:
                run = para.add_run(clean)
                set_run_font(run, size=sz, color=clr)


# ============================================================
# 문서 생성 및 페이지 설정
# ============================================================

def create_document():
    """A4 문서 생성, PDF와 동일한 여백/페이지 번호 설정"""
    doc = Document()

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = FONT_KR
    style.font.size = BODY_SIZE
    fmt = style.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.5

    # 동아시아 폰트
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:eastAsia="{FONT_KR}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_KR)

    # 푸터: 페이지 번호
    add_page_numbers(section)

    return doc


def add_page_numbers(section):
    """푸터에 'N/전체' 형식 페이지 번호 추가"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def add_field(paragraph, field_code):
        run1 = paragraph.add_run()
        fld_begin = parse_xml(f'<w:fldChar xmlns:w="{ns}" w:fldCharType="begin"/>')
        run1._element.append(fld_begin)

        run2 = paragraph.add_run()
        instr = parse_xml(
            f'<w:instrText xmlns:w="{ns}" xml:space="preserve"> {field_code} </w:instrText>'
        )
        run2._element.append(instr)

        run3 = paragraph.add_run()
        fld_end = parse_xml(f'<w:fldChar xmlns:w="{ns}" w:fldCharType="end"/>')
        run3._element.append(fld_end)

    add_field(p, "PAGE")
    run_slash = p.add_run("/")
    set_run_font(run_slash, size=PAGE_NUM_SIZE, color=GRAY)
    add_field(p, "NUMPAGES")

    for run in p.runs:
        set_run_font(run, size=PAGE_NUM_SIZE, color=GRAY)


# ============================================================
# 테이블 데이터 로드
# ============================================================

def load_tables():
    with open(TABLES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

TABLES = None  # 전역에서 lazy 로드


def get_tables():
    global TABLES
    if TABLES is None:
        TABLES = load_tables()
    return TABLES


# ============================================================
# 테이블 렌더링
# ============================================================

def add_table(doc, key):
    """tables.json의 테이블을 Word 표로 삽입"""
    tables = get_tables()
    entry = tables[key]
    data = entry["data"]
    caption = entry.get("caption", "")

    cols = list(data.keys())
    n_rows = len(data[cols[0]])

    # Word 표 생성
    table = doc.add_table(rows=n_rows + 1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 테이블 스타일: 테두리 있는 그리드
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(borders_xml))

    # 헤더 행
    for j, col_name in enumerate(cols):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        set_run_font(run, size=TABLE_SIZE, bold=True)

        # 헤더 배경색
        shading = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:fill="{TABLE_HEADER_BG}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

    # 데이터 행
    for i in range(n_rows):
        for j, col_name in enumerate(cols):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            value = str(data[col_name][i])

            # 마크다운 볼드 처리
            if '**' in value:
                add_formatted_runs(p, value, base_size=TABLE_SIZE)
            else:
                run = p.add_run(value)
                set_run_font(run, size=TABLE_SIZE)

            # 수치 열: 가운데 정렬
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 캡션
    if caption:
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=Pt(2), space_after=Pt(8))
        add_formatted_runs(p, caption, base_size=CAPTION_SIZE, base_color=GRAY)


# ============================================================
# 이미지 삽입
# ============================================================

def add_image(doc, filename, width=None, caption_text=None):
    """이미지를 삽입하고 캡션을 추가"""
    img_path = IMAGES_DIR / filename
    if not img_path.exists():
        p = doc.add_paragraph(f"[이미지 없음: {filename}]")
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    w = width or Mm(150)
    run.add_picture(str(img_path), width=w)

    if caption_text:
        cp = doc.add_paragraph()
        set_paragraph_format(cp, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(2), space_after=Pt(8))
        add_formatted_runs(cp, caption_text, base_size=CAPTION_SIZE, base_color=GRAY)


def extract_network_graph_from_pdf():
    """PDF에서 네트워크 그래프 이미지를 추출"""
    if not PDF_PATH.exists():
        return None
    try:
        import fitz
        doc = fitz.open(str(PDF_PATH))
        # 페이지 10 (index 9)에 네트워크 그래프가 있음
        for page_idx in [9, 10, 8]:
            if page_idx >= doc.page_count:
                continue
            page = doc[page_idx]
            images = page.get_images()
            for img_info in images:
                xref = img_info[0]
                base_img = doc.extract_image(xref)
                w, h = base_img["width"], base_img["height"]
                # 큰 이미지가 네트워크 그래프
                if w > 1500 and h > 1000:
                    TEMP_DIR.mkdir(parents=True, exist_ok=True)
                    out_path = TEMP_DIR / f"network_graph.{base_img['ext']}"
                    with open(out_path, "wb") as f:
                        f.write(base_img["image"])
                    doc.close()
                    return out_path
        doc.close()
    except ImportError:
        pass
    return None


# ============================================================
# Plotly 차트 → 이미지
# ============================================================

def generate_rank_chart():
    """'哲學' 순위 차이 꺾은선 차트 → PNG"""
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    out_path = TEMP_DIR / "rank_chart.png"

    words = ['哲學', '宗敎', '生命', '進化', '眞理']
    rank_1915 = [8, 1, 3, 116, 260]
    rank_1924 = [518, 11, 2, 8, 31]

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        name='1915 철학과 종교', x=words, y=rank_1915,
        mode='lines+markers+text',
        line=dict(color='#4A90A4', width=3),
        marker=dict(size=12, symbol='circle'),
        text=[f'{r}위' for r in rank_1915],
        textposition='top center',
        textfont=dict(size=11, color='#4A90A4'),
    ))
    fig.add_trace(go.Scatter(
        name='1924 인내천요의', x=words, y=rank_1924,
        mode='lines+markers+text',
        line=dict(color='#E07B54', width=3),
        marker=dict(size=12, symbol='diamond'),
        text=[f'{r}위' for r in rank_1924],
        textposition='bottom center',
        textfont=dict(size=11, color='#E07B54'),
    ))
    fig.add_annotation(
        x='哲學', y=260, text="<b>510계단 차이</b>",
        showarrow=True, arrowhead=2, arrowsize=1.5, arrowwidth=2,
        arrowcolor='#D9534F', ax=0, ay=-40,
        font=dict(size=12, color='#D9534F'),
        bgcolor='rgba(255,255,255,0.8)',
        bordercolor='#D9534F', borderwidth=1,
    )
    fig.update_layout(
        title=dict(text="주요 개념어 순위 대조", font=dict(size=16)),
        xaxis_title='개념어', yaxis_title='순위 (낮을수록 빈도 높음)',
        yaxis=dict(autorange='reversed', range=[0, 550]),
        legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='center', x=0.5),
        height=500, width=800, margin=dict(t=80, b=60),
        plot_bgcolor='white',
    )

    fig.write_image(str(out_path), scale=2)
    return out_path


def generate_chapter_dist_chart():
    """장별 참조쌍 분포 막대 차트 → PNG"""
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    out_path = TEMP_DIR / "chapter_dist_chart.png"

    chapters = ['C01 緖言', 'C02 天道', 'C03 眞理', 'C04 目的', 'C05 修煉', 'C06 雜感']
    counts = [3, 0, 45, 0, 0, 63]
    colors = ['#95a5a6', '#95a5a6', '#3498db', '#95a5a6', '#95a5a6', '#e74c3c']

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=chapters, y=counts, marker_color=colors,
        text=counts, textposition='outside',
    ))
    fig.update_layout(
        title=dict(text="1924 텍스트 장별 유효 참조쌍 분포 (총 111개)",
                   font=dict(size=16)),
        yaxis_title='참조쌍 개수',
        height=400, width=800, margin=dict(t=60, b=40),
        plot_bgcolor='white',
    )
    fig.write_image(str(out_path), scale=2)
    return out_path


# ============================================================
# 마크다운 파싱 및 블록 변환
# ============================================================

def parse_content_to_blocks(content):
    """마크다운 콘텐츠를 블록 리스트로 변환

    반환: [('type', 'content'), ...]
    type: heading1, heading2, heading3, paragraph, blockquote,
          list, hr, html_block, footnote_block, component, example_box
    """
    blocks = []
    # 컴포넌트 마커로 분할
    parts = re.split(r'(<!-- component: \w+ -->)', content)

    for part in parts:
        m = re.match(r'<!-- component: (\w+) -->', part)
        if m:
            blocks.append(('component', m.group(1)))
        else:
            blocks.extend(_parse_text_to_blocks(part))
    return blocks


def _parse_text_to_blocks(text):
    """텍스트를 블록으로 분할"""
    blocks = []
    lines = text.split('\n')
    buf = []
    buf_type = None

    def flush():
        nonlocal buf, buf_type
        if buf:
            joined = '\n'.join(buf).strip()
            if joined:
                blocks.append((buf_type, joined))
            buf = []
            buf_type = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄: 현재 버퍼 플러시
        if not stripped:
            flush()
            i += 1
            continue

        # HTML 블록 감지: <div ...>
        if stripped.startswith('<div ') or stripped == '<div>':
            flush()
            # </div>까지 수집
            html_buf = [stripped]
            depth = stripped.count('<div') - stripped.count('</div>')
            i += 1
            while i < len(lines) and depth > 0:
                html_buf.append(lines[i])
                depth += lines[i].count('<div') - lines[i].count('</div>')
                i += 1
            joined = '\n'.join(html_buf).strip()
            # 각주 블록인지 판단
            if 'font-size: 13px' in joined and '<sup>' in joined:
                blocks.append(('footnote_block', joined))
            elif 'border-left:' in joined and 'background' in joined:
                blocks.append(('example_box', joined))
            else:
                blocks.append(('html_block', joined))
            continue

        # 헤딩
        if stripped.startswith('#### '):
            flush()
            blocks.append(('heading3', stripped[5:]))
            i += 1
            continue
        if stripped.startswith('### '):
            flush()
            blocks.append(('heading2', stripped[4:]))
            i += 1
            continue
        if stripped.startswith('## '):
            flush()
            blocks.append(('heading1', stripped[3:]))
            i += 1
            continue

        # 수평선
        if stripped == '---':
            flush()
            blocks.append(('hr', ''))
            i += 1
            continue

        # 인용문 (>)
        if stripped.startswith('> ') or stripped == '>':
            if buf_type != 'blockquote':
                flush()
                buf_type = 'blockquote'
            content = stripped[2:] if stripped.startswith('> ') else ''
            buf.append(content)
            i += 1
            continue

        # 목록 (-)
        if stripped.startswith('- ') and buf_type != 'paragraph':
            if buf_type != 'list':
                flush()
                buf_type = 'list'
            buf.append(stripped[2:])
            i += 1
            continue

        # 일반 문단
        if buf_type != 'paragraph':
            flush()
            buf_type = 'paragraph'
        # 줄 앞의 '+' 제거 (diff 아티팩트)
        if stripped.startswith('+') and not stripped.startswith('++'):
            stripped = stripped[1:]
        buf.append(stripped)
        i += 1

    flush()
    return blocks


# ============================================================
# 블록을 Word 문단으로 변환
# ============================================================

def render_blocks(doc, blocks, components):
    """블록 리스트를 Word 문서에 추가"""
    for block_type, content in blocks:
        if block_type == 'heading1':
            add_heading_paragraph(doc, content, level=1)
        elif block_type == 'heading2':
            add_heading_paragraph(doc, content, level=2)
        elif block_type == 'heading3':
            add_heading_paragraph(doc, content, level=3)
        elif block_type == 'paragraph':
            add_body_paragraph(doc, content)
        elif block_type == 'blockquote':
            add_blockquote(doc, content)
        elif block_type == 'list':
            add_list_items(doc, content)
        elif block_type == 'hr':
            add_toc(doc)  # PDF에서 --- 위치에 목차 박스가 들어감
        elif block_type == 'footnote_block':
            add_footnote_block(doc, content)
        elif block_type == 'example_box':
            add_example_box(doc, content)
        elif block_type == 'html_block':
            add_html_block(doc, content)
        elif block_type == 'component':
            render_component(doc, content, components)


def add_heading_paragraph(doc, text, level=1):
    """장/절/항 제목 추가"""
    p = doc.add_paragraph()
    sizes = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}
    sz = sizes.get(level, H2_SIZE)
    space = {1: Pt(18), 2: Pt(14), 3: Pt(10)}

    set_paragraph_format(p,
                         space_before=space.get(level, Pt(10)),
                         space_after=Pt(6),
                         keep_with_next=True)
    add_formatted_runs(p, text, base_size=sz)
    # 모든 런을 볼드로
    for run in p.runs:
        run.bold = True


def add_body_paragraph(doc, text):
    """본문 문단 추가"""
    # 여러 줄이 합쳐진 경우 공백으로 연결
    text = re.sub(r'\n', ' ', text).strip()
    if not text:
        return
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(2), space_after=Pt(4),
                         line_spacing=1.5)
    add_formatted_runs(p, text, base_size=BODY_SIZE)


def add_blockquote(doc, text):
    """인용문 블록 추가 (들여쓰기 + 작은 글씨)"""
    # 빈 줄로 분리된 단락들
    paras = text.split('\n\n') if '\n\n' in text else [text]
    for para_text in paras:
        para_text = re.sub(r'\n', ' ', para_text).strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p,
                             indent_left=Mm(10),
                             space_before=Pt(2),
                             space_after=Pt(2),
                             line_spacing=1.4)
        add_formatted_runs(p, para_text, base_size=BLOCKQUOTE_SIZE, base_color=DARK_GRAY)


def add_list_items(doc, text):
    """목록 항목 추가"""
    items = text.split('\n')
    for item in items:
        item = item.strip()
        if not item:
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p, indent_left=Mm(8), space_before=Pt(1),
                             space_after=Pt(1), line_spacing=1.4)
        run_bullet = p.add_run("• ")
        set_run_font(run_bullet, size=BODY_SIZE)
        add_formatted_runs(p, item, base_size=BODY_SIZE)


def add_footnote_block(doc, html_content):
    """각주 블록 추가 (작은 글씨, 위쪽 구분선)"""
    # 구분선
    p_line = doc.add_paragraph()
    set_paragraph_format(p_line, space_before=Pt(12), space_after=Pt(4))
    # 얇은 수평선 (문단 아래 테두리)
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # HTML에서 각주 텍스트 추출
    # <sup>N</sup> text<br> 패턴
    text = re.sub(r'<div[^>]*>', '', html_content)
    text = text.replace('</div>', '')
    lines = re.split(r'<br\s*/?>', text)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=Pt(1), space_after=Pt(1),
                             line_spacing=1.3)
        add_formatted_runs(p, line, base_size=FOOTNOTE_SIZE, base_color=GRAY)


def add_example_box(doc, html_content):
    """예시 박스 추가 (왼쪽 테두리가 있는 인용문)"""
    # HTML 태그에서 텍스트 추출
    text = re.sub(r'<div[^>]*>', '', html_content)
    text = text.replace('</div>', '')
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')

    # 색상 감지
    is_1924 = 'ff7f0e' in html_content

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # span 색상 제거하되 텍스트 유지
        line = re.sub(r'<span[^>]*>', '', line).replace('</span>', '')
        p = doc.add_paragraph()
        set_paragraph_format(p, indent_left=Mm(8), space_before=Pt(1),
                             space_after=Pt(1), line_spacing=1.3)
        add_formatted_runs(p, line, base_size=BLOCKQUOTE_SIZE, base_color=DARK_GRAY)

    # 박스 뒤 여백
    p_space = doc.add_paragraph()
    set_paragraph_format(p_space, space_before=Pt(0), space_after=Pt(4))


def add_html_block(doc, html_content):
    """일반 HTML 블록 → 텍스트만 추출하여 추가"""
    # 제목 블록 (text-align: center) 특별 처리
    if 'text-align: center' in html_content:
        render_title_block(doc, html_content)
        return
    if 'text-align: right' in html_content:
        render_author_block(doc, html_content)
        return

    # 기타 HTML: 태그 제거하고 텍스트만
    text = strip_html_tags(html_content)
    text = clean_html_entities(text)
    if text.strip():
        add_body_paragraph(doc, text.strip())


def render_title_block(doc, html_content):
    """제목 블록 렌더링"""
    # h2 내용 추출
    h2_match = re.search(r'<h2[^>]*>(.*?)</h2>', html_content, re.DOTALL)
    h3_match = re.search(r'<h3[^>]*>(.*?)</h3>', html_content, re.DOTALL)

    if h2_match:
        title_text = strip_html_tags(h2_match.group(1)).strip()
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(30), space_after=Pt(4))
        run = p.add_run(title_text)
        set_run_font(run, size=TITLE_SIZE, bold=True)

    if h3_match:
        subtitle_text = strip_html_tags(h3_match.group(1)).strip()
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=Pt(0), space_after=Pt(6))
        run = p.add_run(subtitle_text)
        set_run_font(run, size=SUBTITLE_SIZE, color=GRAY)


def render_author_block(doc, html_content):
    """발표자 블록 렌더링"""
    text = strip_html_tags(html_content).strip()
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=Pt(6), space_after=Pt(12))
    run = p.add_run(text)
    set_run_font(run, size=AUTHOR_SIZE, color=GRAY)


# ============================================================
# 컴포넌트 렌더링 (테이블, 차트, 이미지)
# ============================================================

def render_component(doc, name, components):
    """컴포넌트 마커를 실제 콘텐츠로 변환"""
    handler = components.get(name)
    if handler:
        handler(doc)
    else:
        p = doc.add_paragraph(f"[컴포넌트: {name}]")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0, 0)


# ============================================================
# 목차 (TOC) 생성
# ============================================================

def add_toc(doc):
    """PDF와 동일한 형태의 목차 박스 추가"""
    toc_items = [
        ("Ⅰ. 머리말", True, None),
        ("Ⅱ. DB 구축과 다국어 비교 방법", True, None),
        ("1. 문단별 데이터를 비교 단위로", False, None),
        ("2. 공통 한자어 비율을 계산", False, None),
        ("Ⅲ. 취사선택 양상과 특징", True, None),
        ("1. 참조쌍의 분포: 두 장에 97%가 집중", False, None),
        ("2. 가져온 것과 버린 것", False, None),
        ("Ⅳ. 철학 개념을 다루는 방식", True, None),
        ("1. '철학'이라는 이름 지우기", False, None),
        ("2. 원문에서 읽는 대체 패턴", False, None),
        ("3. 패턴의 검증", False, None),
        ("Ⅴ. 맺음말", True, None),
        ("부록", False, None),
    ]

    # 테두리 있는 1셀 테이블로 목차 박스 구현
    toc_table = doc.add_table(rows=1, cols=1)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = toc_table.rows[0].cells[0]
    cell.text = ""

    # 테두리 설정
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="999999"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

    # 첫 문단 제거 후 항목 추가
    for idx, (text, is_main, _) in enumerate(toc_items):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        sz = TOC_MAIN_SIZE if is_main else TOC_SUB_SIZE
        indent = Mm(0) if is_main else Mm(8)
        set_paragraph_format(p, indent_left=indent,
                             space_before=Pt(2) if is_main else Pt(0),
                             space_after=Pt(0), line_spacing=1.4)
        run = p.add_run(text)
        set_run_font(run, size=sz, bold=is_main)


# ============================================================
# 웹 URL 표시
# ============================================================

def add_web_url_notice(doc):
    """PDF 상단의 웹 URL 공지 추가"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(8))
    run1 = p.add_run("※ 본 발표문의 인터랙티브 웹 버전: ")
    set_run_font(run1, size=WEB_URL_SIZE, color=GRAY)
    run2 = p.add_run(WEB_URL)
    set_run_font(run2, size=WEB_URL_SIZE, color=BLUE_LINK)


# ============================================================
# 섹션 처리 함수
# ============================================================

def load_section(filename):
    """sections/ 폴더에서 마크다운 파일 로드"""
    path = SECTIONS_DIR / filename
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def process_section(doc, filename, components):
    """마크다운 섹션 파일을 파싱하여 Word에 추가"""
    content = load_section(filename)
    blocks = parse_content_to_blocks(content)
    render_blocks(doc, blocks, components)


# ============================================================
# 섹션별 컴포넌트 매핑
# ============================================================

def build_components():
    """모든 컴포넌트 핸들러를 생성"""

    # --- 테이블 핸들러 ---
    def make_table_handler(key):
        return lambda doc: add_table(doc, key)

    # --- 이미지/차트 핸들러 ---
    def heatmap_handler(doc):
        add_image(doc, "chapter_heatmap_avg_above_threshold.png",
                  width=Mm(120),
                  caption_text="111개 유효 참조쌍 (자카드 유사도 ≥ 0.1). "
                               "색상은 참조쌍 개수, 괄호 안은 평균 유사도.")

    def network_graph_handler(doc):
        img_path = extract_network_graph_from_pdf()
        if img_path and img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Mm(155))
            # 캡션
            cp = doc.add_paragraph()
            set_paragraph_format(cp, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 space_before=Pt(2), space_after=Pt(8))
            run_c = cp.add_run("111개 유효 참조쌍의 문단 네트워크. "
                               "왼쪽 ● = 1915 문단 52개, 오른쪽 ◆ = 1924 문단 35개.")
            set_run_font(run_c, size=CAPTION_SIZE, color=GRAY)
        else:
            p = doc.add_paragraph("[네트워크 그래프 - 웹 버전 참조]")
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            p.runs[0].font.color.rgb = GRAY

    def paradox_image_handler(doc):
        add_image(doc, "c03s04_philosophy_paradox.png",
                  width=Mm(150),
                  caption_text="C03-S04 항별 참조쌍 개수(막대)와 '哲學' 출현 빈도(선). "
                               "참조 밀도가 높은 I02, I05에서 '哲學'은 0회.")

    def rank_chart_handler(doc):
        try:
            img_path = generate_rank_chart()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Mm(140))
        except Exception as e:
            p = doc.add_paragraph(f"[순위 차트 생성 실패: {e}]")

    def chapter_dist_chart_handler(doc):
        try:
            img_path = generate_chapter_dist_chart()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Mm(140))
        except Exception as e:
            p = doc.add_paragraph(f"[분포 차트 생성 실패: {e}]")

    def pairs_data_handler(doc):
        csv_path = PROJECT_ROOT / "data" / "analysis" / "validated_pairs_final.csv"
        if csv_path.exists():
            df = pd.read_csv(csv_path)
            cols = ['rank', 'similarity', 'pid_1915', 'pid_1924', 'common_tokens']
            available = [c for c in cols if c in df.columns]
            if available:
                df_show = df[available].head(20)
                table = doc.add_table(rows=len(df_show) + 1, cols=len(available))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 헤더
                col_labels = {
                    'rank': '순위', 'similarity': '유사도',
                    'pid_1915': '1915 문단', 'pid_1924': '1924 문단',
                    'common_tokens': '공통 토큰'
                }
                for j, col in enumerate(available):
                    cell = table.rows[0].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(col_labels.get(col, col))
                    set_run_font(run, size=Pt(8), bold=True)
                    shading = parse_xml(
                        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        f'w:fill="{TABLE_HEADER_BG}" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)
                # 데이터
                for i, (_, row) in enumerate(df_show.iterrows()):
                    for j, col in enumerate(available):
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        val = str(row[col])
                        if col == 'similarity':
                            val = f"{row[col]:.4f}"
                        run = p.add_run(val)
                        set_run_font(run, size=Pt(8))
                        if j < 2:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cp = doc.add_paragraph()
                run_c = cp.add_run("*전체 111개 참조쌍은 data/analysis/validated_pairs_final.csv 참조*")
                set_run_font(run_c, size=CAPTION_SIZE, italic=True, color=GRAY)
        else:
            add_table(doc, "pairs_fallback_table")

    # 컴포넌트 매핑
    return {
        # Ⅱ-1
        "corpus_table": make_table_handler("corpus_table"),
        "paragraph_table": make_table_handler("paragraph_table"),
        # Ⅱ-2
        "stats_table": make_table_handler("stats_table"),
        "common_tokens_table": make_table_handler("common_tokens_table"),
        "range_example_table": make_table_handler("range_example_table"),
        "valid_pairs_table": make_table_handler("valid_pairs_table"),
        # Ⅲ-1
        "chapter_dist_table": make_table_handler("chapter_dist_table"),
        "heatmap": heatmap_handler,
        "network_graph": network_graph_handler,
        "ref_table": make_table_handler("ref_table"),
        "c03s04_table": make_table_handler("c03s04_table"),
        "c06s06_table": make_table_handler("c06s06_table"),
        # Ⅲ-2
        "borrow_table": make_table_handler("borrow_table"),
        "invert_table": make_table_handler("invert_table"),
        # Ⅳ-1
        "phil_freq_table": make_table_handler("phil_freq_table"),
        "freq_compare_table": make_table_handler("freq_compare_table"),
        "inverse_table": make_table_handler("inverse_table"),
        # Ⅳ-2
        "sub_stats_table": make_table_handler("sub_stats_table"),
        "pattern_table": make_table_handler("pattern_table"),
        # Ⅳ-3
        "diff_vocab_table": make_table_handler("diff_vocab_table"),
        "phil_compound_table": make_table_handler("phil_compound_table"),
        # 부록
        "file_table": make_table_handler("file_table"),
        "pairs_data": pairs_data_handler,
    }


# ============================================================
# 메인 함수
# ============================================================

def main():
    print("=" * 50)
    print("발표문 Word 변환 시작")
    print("=" * 50)

    doc = create_document()
    components = build_components()

    # ── 웹 URL 공지 ──
    add_web_url_notice(doc)

    # ── Ⅰ. 머리말 ──
    print("[1/10] Ⅰ. 머리말...")
    process_section(doc, "01_머리말.md", components)

    # ── 목차 삽입 (머리말 제목 앞 대신, 여기서는 제목 뒤에) ──
    # PDF에서는 제목+부제+발표자 → 목차 → Ⅰ.머리말 순서
    # 01_머리말.md에 이미 제목+Ⅰ.머리말이 있으므로,
    # 목차는 제목 블록 바로 뒤에 삽입해야 하지만,
    # 간단히 Ⅰ장 뒤에 넣지 않고, 빌드 순서를 조정

    # ── Ⅱ. DB 구축과 다국어 비교 방법 ──
    print("[2/10] Ⅱ. DB 구축과 다국어 비교 방법...")
    add_heading_paragraph(doc, "Ⅱ. DB 구축과 다국어 비교 방법", level=1)
    process_section(doc, "02-1_문단별_데이터.md", components)
    process_section(doc, "02-2_공통_한자어_비율.md", components)

    # ── Ⅲ. 취사선택 양상과 특징 ──
    print("[3/10] Ⅲ. 취사선택 양상과 특징...")
    add_heading_paragraph(doc, "Ⅲ. 취사선택 양상과 특징", level=1)
    process_section(doc, "03-1_분포와_특징.md", components)
    process_section(doc, "03-2_보편화와_비교.md", components)

    # ── Ⅳ. 철학 개념을 다루는 방식 ──
    print("[4/10] Ⅳ. 철학 개념을 다루는 방식...")
    add_heading_paragraph(doc, "Ⅳ. 철학 개념을 다루는 방식", level=1)
    process_section(doc, "04-1_철학_기표_소거.md", components)
    process_section(doc, "04-2_철리_인내천.md", components)
    process_section(doc, "04-3_패턴의_검증.md", components)

    # ── Ⅴ. 맺음말 ──
    print("[5/10] Ⅴ. 맺음말...")
    process_section(doc, "05_맺음말.md", components)

    # ── 부록 ──
    print("[6/10] 부록...")
    add_heading_paragraph(doc, "부록", level=1)
    process_section(doc, "06_부록.md", components)

    # ── 저장 ──
    print(f"\n[저장] {OUTPUT_PATH}")
    doc.save(str(OUTPUT_PATH))
    size_kb = os.path.getsize(str(OUTPUT_PATH)) / 1024
    print(f"완료! ({size_kb:.0f} KB)")

    # 임시 파일 정리
    if TEMP_DIR.exists():
        import shutil
        shutil.rmtree(TEMP_DIR, ignore_errors=True)
        print("임시 파일 정리 완료.")


if __name__ == "__main__":
    main()
